from typing import List, Optional
from datetime import datetime
import os
import json
import re
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.table import Table

from document import MedicaDoc

TEMPLATE_FILE = "Template-Premed.docx"

ANAMNESE_ENTS = ['HD', 'ND', 'Allergien']
MED_KEY = "Dauermedikation"

def create_premed_report(entity_dict: dict, pat_name: str) -> Document:
    # load template
    doc = Document(TEMPLATE_FILE)
    # get document table for anamnesis
    table = _get_table_by_name(doc, 'Anamnese')
    if table:
        table_content = table.rows[0].cells[0].paragraphs[1].text
        
        index_dict = {k: table_content.find(k) for k in entity_dict.keys() if k in ANAMNESE_ENTS}

        new_content_len = 0
        # iterate over sorted dict
        for ent_key, key_index in sorted(index_dict.items(), key=lambda item: item[1]):
        #for ent_key, key_index in index_dict.items():
            # build starting index to insert entities
            start_indx = key_index + len(ent_key) + len(':\n') + new_content_len
            # add entities to string
            table_content = table_content[:start_indx] + ','.join(entity_dict[ent_key]) + table_content[start_indx:]
            # increase index
            new_content_len += len(','.join(entity_dict[ent_key]))
        
        # insert new string to document table
        table.rows[0].cells[0].paragraphs[1].text = table_content

    else:
        print("WARNING: cannot find Anamnese table in file")

    meds = entity_dict.get(MED_KEY)
    
    if meds and table.rows[0].cells[1].text.replace('\n', '').strip() == MED_KEY:
        table_content = table.rows[0].cells[1].paragraphs[1].text
        # insert new string to document table
        table.rows[0].cells[1].paragraphs[1].text = table_content + "\n".join(meds)

    pat_info_table = _get_table_by_name(doc, "Patienteninformationen")
    pat_info_table.rows[1].cells[2].paragraphs[1].text = pat_name

    # CAVE
    if entity_dict['cave']:
        cave_table = _get_table_by_name(doc, "Diagnose")
        cave_table.rows[0].cells[0].paragraphs[1].text = ', '.join(entity_dict['cave'])
    
    return doc

def premed_dict_to_markdown(premed_dict: dict) -> str:
    mkrd_str = ""
    if premed_dict:
        for k in ['HD', 'ND', 'Dauermedikation', 'Allergien', 'Behandlung']:
            mkrd_str += f"#### {k} \n"
            if premed_dict.get(k):
                for ele in premed_dict[k]:
                    mkrd_str += f"{ele}  \n"
    return mkrd_str


def _get_table_by_name(doc: Document, table_name: str) -> Optional[Table]:
    for table in doc.tables:
        if table.rows[0].cells[0].text.split()[0].lower() == table_name.lower():
            return table


def highlight_words(section_entities: List[dict], color: str = 'blue') -> str:
    sentence = section_entities[0]['sentence'].strip()
    for entity in section_entities:
        sentence = color_word(sentence, entity['word'], color)
    return sentence

def color_word(sentence: str, word: str, color: str = 'blue') -> str:
    idx = sentence.find(word)
    # in case match was not found
    if idx == -1:
        word = remove_spaces_except_hyphen(word)
        idx = sentence.find(word)
        if idx != -1:
            sentence = sentence[:idx] + f"**:{color}[{word}]**" + sentence[idx + len(word):]    
    else:
        sentence = sentence[:idx] + f"**:{color}[{word}]**" + sentence[idx + len(word):]
    return sentence


def remove_spaces_except_hyphen(input_string):
    # Define the pattern for splitting by '-:/'
    split_pattern = r'([-:/])'

    # Split the input string based on the pattern
    parts = re.split(split_pattern, input_string)

    # Remove spaces from non-hyphen parts
    cleaned_parts = [part.strip() if part not in '-:/' else part for part in parts]

    # Join the parts
    result_string = ''.join(cleaned_parts)

    return result_string


def aggregate_entity_info(entities: List[dict]) -> dict:
    """Aggregates entities over multiple documents"""
    agg_entities: dict = {}
    for i, entity in enumerate(entities):

        if entity['word'] not in agg_entities:
            # init info
            agg_entities[entity['word']] = {
                'word': entity['word'],
                'type': [entity['entity_group']],
                'source_indices': [i],
                'first_occurance': entity['date'],
                'last_occurance': None,
                'negated': entity['negated'],
                'docs': [entity['doc']]
            }
        else:
            # add info
            # type
            if entity['entity_group'] not in agg_entities[entity['word']]['type']:
                agg_entities[entity['word']]['type'].append(entity['entity_group'])

            # source index
            agg_entities[entity['word']]['source_indices'].append(i)

            # doc index, we always add doc index to count occurances
            agg_entities[entity['word']]['docs'].append(entity['doc'])

            # first, last occurance
            first_dt = datetime.strptime(agg_entities[entity['word']]['first_occurance'], '%Y-%m-%d')
            ent_dt = datetime.strptime(entity['date'], '%Y-%m-%d')
            # new first?
            if ent_dt < first_dt:
                agg_entities[entity['word']]['first_occurance'] = entity['date']
                # set last date
                if agg_entities[entity['word']]['last_occurance'] is None:
                    agg_entities[entity['word']]['last_occurance'] = first_dt.strftime('%Y-%m-%d')
            # new end date
            elif agg_entities[entity['word']]['last_occurance'] is None:
                agg_entities[entity['word']]['last_occurance'] = entity['date']
            elif ent_dt > datetime.strptime(agg_entities[entity['word']]['last_occurance'], '%Y-%m-%d'):
                agg_entities[entity['word']]['last_occurance'] = entity['date']
            
            # negation
            # we only set flag of negation if all entites are negated
            if entity['negated'] is False and agg_entities[entity['word']]['negated'] is True:
                agg_entities[entity['word']]['negated'] = False

    return agg_entities


def timeline_plot(dates):
    """Returns plot figure of document dates"""
    sorted_dates = sorted(dates)
    fig, ax = plt.subplots()
    fig.set_figheight(1)
    fig.set_figwidth(12)
    ax.plot(
        sorted_dates,
        [1 for i in range(len(dates))],
        'o:',
        markersize=12,
        color='#F63366'
    )
    ax.get_yaxis().set_visible(False)
    ax.spines[['right', 'top', 'left']].set_visible(False)
    return fig


def exclude_types(df: pd.DataFrame, types: List[str]) -> pd.DataFrame:
    """Method to filter dataframe"""
    def _includes_types(row):
        for t in types:
            if t in row['type']:
                return True
        return False
    not_empty = False
    if df.empty:
        pass
    else:
        not_empty = True
    
    df['filter'] = df.apply(lambda x: _includes_types(x), axis=1)

    if df[df['filter'] == False].empty and not_empty:
        print("HERE we go :/")

    return df[df['filter'] == False]


def read_config(file_path: str) -> dict:
    with open(file_path, 'r') as fh:
        cfg = json.load(fh)
    return cfg


def store_doc(subject: str, doc_date: str, medica_doc: MedicaDoc):
    # check if name exists
    file_path = f"data/{subject}/{subject}_{doc_date}.json"
    if os.path.isfile(file_path):
        # add incremental number
        file_number = get_file_number(file_path)
        if file_number:
            file_path = f"data/{subject}/{subject}_{doc_date}({file_number}).json"
    
    # check if directory exists 
    if not os.path.isdir(os.path.dirname(file_path)):
        os.mkdir(os.path.dirname(file_path))
    # write json file        
    with open(file_path, "w") as fh:
        json.dump(medica_doc.to_dict(), fh)


def get_file_number(file_path: str) -> Optional[str]:
    """Method to check for documents with same date, if so increase count"""
    file_name, ext = os.path.splitext(os.path.basename(file_path))
    # loop over files
    same_file_name_numbers = []
    for file in os.listdir(os.path.dirname(file_path)):
        f_name = os.path.splitext(file)[0]
        if f_name.split('(')[0] == file_name:
            number_match = re.search(r'\([0-9]*\)', f_name)
            if number_match:
                same_file_name_numbers.append(int(number_match.group()[1:-1]))
            else:
                same_file_name_numbers.append(0)
    if same_file_name_numbers:
        return str(max(same_file_name_numbers) + 1)
